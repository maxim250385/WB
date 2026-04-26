# =============================================================================
#  modules/file_reader.py — Чтение файлов всех поддерживаемых форматов
# =============================================================================

from __future__ import annotations
import os
import base64
from pathlib import Path

from config import (
    TEXT_EXTENSIONS, OFFICE_EXTENSIONS, PDF_EXTENSIONS,
    IMAGE_EXTENSIONS, TABLE_EXTENSIONS, ALL_SUPPORTED,
)


def read_file(path: str) -> tuple[str, list[bytes]]:
    """
    Читает файл. Возвращает (text, image_bytes_list).
    """
    ext = Path(path).suffix.lower()

    if ext in TEXT_EXTENSIONS:
        return _read_text(path), []
    if ext in PDF_EXTENSIONS:
        return _read_pdf(path), []
    if ext in OFFICE_EXTENSIONS:
        return _read_docx(path), []
    if ext in IMAGE_EXTENSIONS:
        return "", [_read_image_bytes(path)]
    if ext in TABLE_EXTENSIONS:
        return _read_table(path), []

    print(f"  [!] Неизвестный формат: {ext} — пропускаю {os.path.basename(path)}")
    return "", []


def _read_text(path: str) -> str:
    try:
        with open(path, "r", encoding="utf-8", errors="replace") as f:
            return f.read().strip()
    except Exception as e:
        print(f"  [!] Ошибка чтения текста {path}: {e}")
        return ""


def _read_pdf(path: str) -> str:
    # Попытка 1: pdfplumber (лучше с таблицами)
    try:
        import pdfplumber
        pages_text = []
        with pdfplumber.open(path) as pdf:
            for page in pdf.pages:
                tables = page.extract_tables()
                if tables:
                    for table in tables:
                        for row in table:
                            row_clean = [str(c).strip() for c in row if c is not None]
                            if any(row_clean):
                                pages_text.append(" | ".join(row_clean))
                text = page.extract_text()
                if text:
                    pages_text.append(text.strip())
        result = "\n".join(pages_text).strip()
        if result:
            return result
    except ImportError:
        pass
    except Exception as e:
        print(f"  [!] pdfplumber ошибка: {e}")

    # Попытка 2: pypdf
    try:
        from pypdf import PdfReader
        reader = PdfReader(path)
        parts  = [page.extract_text().strip() for page in reader.pages
                  if page.extract_text()]
        return "\n".join(parts).strip()
    except ImportError:
        print("  [!] Установи: pip install pdfplumber pypdf")
        return ""
    except Exception as e:
        print(f"  [!] Ошибка чтения PDF: {e}")
        return ""


def _read_docx(path: str) -> str:
    ext = Path(path).suffix.lower()
    if ext == ".doc":
        try:
            import docx2txt
            return docx2txt.process(path).strip()
        except ImportError:
            print("  [!] Установи: pip install docx2txt")
            return ""
        except Exception as e:
            print(f"  [!] Ошибка .doc: {e}")
            return ""

    try:
        from docx import Document
        doc   = Document(path)
        parts = []
        for para in doc.paragraphs:
            t = para.text.strip()
            if t:
                parts.append(t)
        for table in doc.tables:
            for row in table.rows:
                row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if row_text:
                    parts.append(" | ".join(row_text))
        return "\n".join(parts).strip()
    except ImportError:
        print("  [!] Установи: pip install python-docx")
        return ""
    except Exception as e:
        print(f"  [!] Ошибка docx: {e}")
        return ""


def _read_image_bytes(path: str) -> bytes:
    try:
        with open(path, "rb") as f:
            return f.read()
    except Exception as e:
        print(f"  [!] Ошибка чтения изображения: {e}")
        return b""


def image_to_base64(image_bytes: bytes) -> str:
    return base64.b64encode(image_bytes).decode("utf-8")


def get_image_mime(path: str) -> str:
    ext = Path(path).suffix.lower()
    return {".jpg": "image/jpeg", ".jpeg": "image/jpeg", ".png": "image/png",
            ".webp": "image/webp", ".bmp": "image/bmp"}.get(ext, "image/jpeg")


def _read_table(path: str) -> str:
    ext = Path(path).suffix.lower()
    try:
        import pandas as pd
        if ext == ".csv":
            for enc in ("utf-8", "cp1251", "latin-1"):
                try:
                    df = pd.read_csv(path, encoding=enc, nrows=200)
                    break
                except UnicodeDecodeError:
                    continue
            else:
                return ""
        else:
            df = pd.read_excel(path, nrows=200)
        rows = [" | ".join(str(c) for c in df.columns)]
        for _, row in df.iterrows():
            rows.append(" | ".join(str(v) for v in row.values))
        return "\n".join(rows)
    except ImportError:
        print("  [!] Установи: pip install pandas openpyxl")
        return ""
    except Exception as e:
        print(f"  [!] Ошибка таблицы: {e}")
        return ""


def read_product_folder(folder_path: str) -> tuple[str, list[tuple[bytes, str]]]:
    """Читает все поддерживаемые файлы из папки товара."""
    all_texts : list[str]               = []
    all_images: list[tuple[bytes, str]] = []

    files = sorted(os.listdir(folder_path))
    for fname in files:
        fpath = os.path.join(folder_path, fname)
        if not os.path.isfile(fpath):
            continue
        ext = Path(fname).suffix.lower()
        if ext not in ALL_SUPPORTED:
            continue

        print(f"    [→] Читаю: {fname}")
        text, image_bytes_list = read_file(fpath)

        if text:
            all_texts.append(f"--- {fname} ---\n{text}")
        for img_bytes in image_bytes_list:
            if img_bytes:
                mime = get_image_mime(fpath)
                all_images.append((img_bytes, mime))

    return "\n\n".join(all_texts), all_images
